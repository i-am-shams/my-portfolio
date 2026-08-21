"""
Generates the Word resume from data/resume.json, so the .docx, the site's /cv page
and public/resume.txt all say the same thing. Before this, the .docx was maintained
by hand and had drifted: it still claimed "two self-built systems" when there were
three, listed two projects, and carried the older management-first experience
bullets that data/resume.json had already moved past.

Usage:
    python scripts/build_resume_docx.py [--photo path/to/photo.png] [-o OUTPUT]

Deliberately writes to a NEW file rather than overwriting an existing draft.

On the photo: omitted by default. For US/UK/Canada/Australia applications a
photo on a CV is a liability - many employers must discard photo CVs for
discrimination-compliance reasons, and some ATS handle the embedded image badly.
Pass --photo only for a local Dhaka variant where it is conventional. It is
downscaled to 220px wide regardless; the original draft embedded a 3024x4032
phone photo, which was 11.5MB of a 12MB file and the reason the exported PDF
was 2MB.
"""

import argparse
import json
import pathlib
import re
import sys
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
CAREER_START = date(2010, 4, 1)

# A web page can carry every bullet; a resume cannot. data/resume.json is ordered
# best-first (see the engineering-first pass in HANDOVER Session 8), so trimming
# from the end drops the weakest lines rather than arbitrary ones.
MAX_BULLETS_RECENT = 5
MAX_BULLETS_OLDER = 3
RECENT_ROLES = 2

INK = RGBColor(0x0F, 0x17, 0x2A)      # slate-900, matches the site
ACCENT = RGBColor(0x1D, 0x4E, 0xD8)   # blue-700, the site's single accent
MUTED = RGBColor(0x47, 0x55, 0x69)    # slate-600

CONTACT = {
    "phone": "+88 01674592829",
    "email": "i.am.shams@gmail.com",
    "site": "khalid-shams.vercel.app",
    "github": "github.com/i-am-shams",
    "availability": "Dhaka, Bangladesh (UTC+6)  |  Open to remote roles worldwide or on-site in Dhaka  |  One month notice",
}


def years_of_experience(today=None):
    today = today or date.today()
    years = today.year - CAREER_START.year
    if (today.month, today.day) < (CAREER_START.month, CAREER_START.day):
        years -= 1
    return years


def set_margins(doc, inches=0.6):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def style_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.02
    # Calibri for East Asian runs too, so Word does not substitute a fallback.
    rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
    rpr.set(qn("w:eastAsia"), "Calibri")


def para(doc, text="", size=9.5, bold=False, italic=False, color=INK,
         space_before=0, space_after=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color
    return p


def bottom_rule(paragraph, color="1D4ED8", size=6):
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(pbdr)


def section_heading(doc, text):
    p = para(doc, text.upper(), size=10.5, bold=True, color=ACCENT,
             space_before=9, space_after=3)
    p.runs[0].font.name = "Calibri"
    bottom_rule(p)
    return p


def bullet(doc, text, indent=0.16):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.02
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = INK
    return p


def role_header(doc, left, right):
    """Position/company on the left, dates right-aligned on the same line.

    A missing `right` must not emit the tab, or Word renders a dangling tab stop
    where the date should be - which is exactly what the education entries did,
    since data/resume.json has no years for them yet.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(7.3), WD_ALIGN_PARAGRAPH.RIGHT)
    a = p.add_run(left)
    a.bold = True
    a.font.size = Pt(10)
    a.font.color.rgb = INK
    if right:
        b = p.add_run("\t" + right)
        b.italic = True
        b.font.size = Pt(9)
        b.font.color.rgb = MUTED
    return p


def add_photo(doc, photo_path):
    from PIL import Image

    img = Image.open(photo_path).convert("RGB")
    img.thumbnail((220, 220))
    tmp = ROOT / "scripts" / "_resume_photo.jpg"
    img.save(tmp, "JPEG", quality=82, optimize=True)
    doc.add_picture(str(tmp), width=Inches(1.05))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tmp.unlink(missing_ok=True)


def build(resume, out_path, photo=None):
    doc = Document()
    set_margins(doc)
    style_base(doc)

    if photo:
        add_photo(doc, photo)

    # Header
    name = para(doc, resume["personalInfo"]["name"].upper(), size=21, bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER)
    name.runs[0].font.color.rgb = INK
    para(doc, resume["personalInfo"]["title"], size=11.5, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, CONTACT["availability"], size=8.5, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, f'{CONTACT["phone"]}  |  {CONTACT["email"]}  |  {CONTACT["site"]}  |  {CONTACT["github"]}',
         size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)

    # Summary
    section_heading(doc, "Professional Summary")
    summary = resume["careerObjective"].replace("{{YEARS}}", str(years_of_experience()))
    para(doc, summary, size=9.5, space_after=2)

    # Projects before experience: this candidate's last decade carries
    # delivery-management titles, so the systems he built alone are the evidence
    # that has to land first. See docs/recruiter-review.md.
    section_heading(doc, "Selected Projects")
    for project in resume["projects"]:
        link = project["url"]
        if project.get("repo"):
            link += "  |  " + project["repo"].replace("https://", "")
        role_header(doc, f'{project["name"]} — {project["tagline"]}',
                    link.replace("https://", ""))
        for line in project["bullets"]:
            bullet(doc, line)
        para(doc, project["stack"], size=8.5, italic=True, color=MUTED, space_after=2)

    # Experience
    section_heading(doc, "Professional Experience")
    for index, role in enumerate(resume["experience"]):
        role_header(doc, f'{role["position"]} — {role["company"]}', role["duration"])
        cap = MAX_BULLETS_RECENT if index < RECENT_ROLES else MAX_BULLETS_OLDER
        for line in role["responsibilities"][:cap]:
            bullet(doc, line)

    # Skills
    section_heading(doc, "Technical Skills")
    for group in resume["technicalSkills"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.02
        label = p.add_run(f'{group["category"]}: ')
        label.bold = True
        label.font.size = Pt(9)
        label.font.color.rgb = INK
        items = p.add_run(" - ".join(group["items"]))
        items.font.size = Pt(9)
        items.font.color.rgb = INK

    # Education and certifications
    section_heading(doc, "Education")
    for entry in resume["education"]:
        right = entry.get("year") or ""
        role_header(doc, f'{entry["degree"]} — {entry["institution"]}', right)

    section_heading(doc, "Certifications")
    for cert in resume["certifications"]:
        para(doc, f'{cert["name"]} — {cert["issuer"]}  |  {cert["validity"]}', size=9.5)

    langs = resume.get("languages") or []
    if langs:
        section_heading(doc, "Languages")
        para(doc, "  |  ".join(f'{l["language"]} ({l["proficiency"]})' for l in langs), size=9.5)

    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("-o", "--output", default="Khalid_Shams_Resume_v2.docx")
    ap.add_argument("--photo", default=None,
                    help="optional headshot; downscaled to 220px. Omit for international applications.")
    args = ap.parse_args()

    resume = json.loads((ROOT / "data" / "resume.json").read_text(encoding="utf-8"))
    if "projects" not in resume:
        sys.exit("data/resume.json has no `projects` key - the resume would omit them silently.")

    out = build(resume, ROOT / args.output, photo=args.photo)
    size_kb = out.stat().st_size / 1024
    print(f"{out.name} written ({size_kb:.0f} KB)")
    if size_kb > 500:
        print("WARNING: over 500 KB - several ATS reject larger uploads.")


if __name__ == "__main__":
    main()
